import logging
import re
import os
import base64
import tempfile
import uuid
import requests
from typing import Tuple
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)


class WordGenerator:
    @staticmethod
    def clean_filename(name: str) -> str:
        """清理文件名中的非法字符"""
        illegal_chars = r'[\/\\:*?"<>|]'
        return re.sub(illegal_chars, "", name)

    @staticmethod
    def extract_images_from_html(html_body: str) -> list[Tuple[str, str]]:
        """从 HTML 中提取图片，返回 [(temp_path, img_format)]"""
        images = []
        if not html_body:
            return images

        soup = BeautifulSoup(html_body, "html.parser")
        img_tags = soup.find_all("img")

        for img in img_tags:
            src = img.get("src", "")
            img_path = None
            fmt = "png"  # 默认格式

            if src.startswith("data:image"):
                # base64 图片
                match = re.match(r"data:image/(\w+);base64,(.+)", src)
                if match:
                    fmt, data = match.groups()
                    try:
                        img_data = base64.b64decode(data)
                        img_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4()}.{fmt}")
                        with open(img_path, "wb") as f:
                            f.write(img_data)
                        logger.info(f"Extracted base64 image to {img_path}")
                    except Exception as e:
                        logger.warning(f"Failed to decode base64 image: {e}")

            elif src.startswith("http"):
                # 网络图片 - 下载
                try:
                    response = requests.get(src, timeout=10)
                    if response.status_code == 200:
                        if "." in src.split("/")[-1]:
                            fmt = src.split("/")[-1].split(".")[-1].lower()
                            if fmt not in ["png", "jpg", "jpeg", "gif", "bmp"]:
                                fmt = "png"
                        img_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4()}.{fmt}")
                        with open(img_path, "wb") as f:
                            f.write(response.content)
                        logger.info(f"Downloaded image from {src} to {img_path}")
                except Exception as e:
                    logger.warning(f"Failed to download image {src}: {e}")

            if img_path and os.path.exists(img_path):
                images.append((img_path, fmt))

        return images

    @staticmethod
    def generate_word(email, output_dir: str) -> str:
        """生成 Word 文档，返回文件路径"""
        doc = Document()

        # 标题
        title = WordGenerator.clean_filename(email.conversation_topic)
        doc.add_heading(title, level=1)

        # 正文（优先使用纯文本，处理 HTML 中的图片）
        body = email.body
        if not body and email.html_body:
            # 从 HTML 提取纯文本
            soup = BeautifulSoup(email.html_body, "html.parser")
            body = soup.get_text()

        doc.add_paragraph(body)

        # 嵌入图片
        if email.html_body:
            images = WordGenerator.extract_images_from_html(email.html_body)
            for img_path, fmt in images:
                try:
                    doc.add_picture(img_path, width=Inches(4))
                except Exception as e:
                    logger.warning(f"Failed to add picture {img_path}: {e}")
                finally:
                    if os.path.exists(img_path):
                        os.remove(img_path)

        # 保存
        os.makedirs(output_dir, exist_ok=True)
        file_path = os.path.join(output_dir, f"{title}.docx")
        doc.save(file_path)
        logger.info(f"Generated Word: {file_path}")
        return file_path
